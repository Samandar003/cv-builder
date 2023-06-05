from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .models import ResumeModel

def cv_build(data, photo):
    doc = Document()

    # Add picture in the top-right corner
    doc.add_picture(photo, width=Inches(2), height=Inches(2))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Save the document
    # Add heading
    doc.add_heading('Curriculum Vitae', level=1)

    # Add personal details
    doc.add_heading('Personal Details', level=2)
    doc.add_paragraph('Name: '+data['name'])
    doc.add_paragraph('Address: '+data['address'])
    doc.add_paragraph('Phone: '+data['phone'])
    doc.add_paragraph('Email: '+data['email'])

    # Add links

    doc.add_heading('Contact Links', level=2)
    doc.add_paragraph('Telegram:  '+data['telegram'])
    doc.add_paragraph('Github:  ' +data['github'])
    doc.add_paragraph('Linkedin:  '+data['linkedin'])
    doc.add_paragraph('Website:  '+data['website'])


    # Add education details
    doc.add_heading('Education', level=2)
    doc.add_paragraph(f"{data['course']}, {data['edu_place']}, {data['start_date']} - {data['end_date']}")

    # Add work experience details
    doc.add_heading('Work Experience', level=2)

    doc.add_paragraph(f"{data['job_role']}, {data['work_place']}, {data['start_date']} - {data['end_date']}")
    doc.add_paragraph(data['job_desc'])



    # Add skill details
    doc.add_paragraph(data['skills'])
        
    # projects

    doc.add_heading('Projects', level=2)
    doc.add_paragraph(data['project1'])
    doc.add_paragraph(data['project2'])
    doc.add_paragraph(data.get('project3', ''))
    doc.add_paragraph(data.get('project4', ''))

    # Certificates
    doc.add_heading('Certificates', level=2)
    doc.add_paragraph(data['certificate1'])
    doc.add_paragraph(data.get('certificate2', ''))

    # Hobbies
    # doc.add_heading('Hobbies', level=2)
    

    # Save the document
    doc.save(f"media/files/{data['name'].title()}_CV.docx")
    outfile = ResumeModel.objects.create(file=f"files/{data['name'].title()}_CV.docx")
    outfile.save()
    
    return {"file":outfile.file}