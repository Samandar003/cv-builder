from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


doc = Document()

# Add picture in the top-right corner
doc.add_picture('men_004.jpg', width=Inches(2), height=Inches(2))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Save the document
# Add heading
doc.add_heading('Curriculum Vitae', level=1)

# Add personal details
doc.add_heading('Personal Details', level=2)
name = input("Name: ")
doc.add_paragraph('Name: '+name)
address = input("Address: ")
doc.add_paragraph('Address: '+address)
phone_num = input("Phone Number: ")
doc.add_paragraph('Phone: '+phone_num)
email = input("email: ")
doc.add_paragraph('Email: '+email)

# num = input("how many links do you wanna input: ")
# for x in range(num):

# link = ""
# Add education details
edu_place = input("Education place: ")
start_date = input("Start_date: ")
end_date = input("End date: ")
course = input("Course: ")
doc.add_heading('Education', level=2)
doc.add_paragraph(f"{course}, {edu_place}, {start_date} - {end_date}")

# Add work experience details
doc.add_heading('Work Experience', level=2)

job_role = input("Your role in the previous job:")
work_place = input("Work Place: ")
start_date = input('Start_date: ')
end_date = input("End date: ")
text = input("Job description: ")

doc.add_paragraph(f'{job_role}, {work_place}, {start_date} - {end_date}')
doc.add_paragraph(text)



# Add skill details
doc.add_heading('Skills', level=2)
num_skills = int(input("How many skills you wanna add: "))
skills_list = ""
for x in range(num_skills):
    skill = input(f"{x+1}-skill: ")
    skills_list += skill + "  "
doc.add_paragraph(skills_list)
    
# projects

doc.add_heading('Projects', level=2)
num_projects = int(input("How many Projects you wanna show: "))
for x in range(num_projects):
    project_title = input(f"{x+1}-project title: ")
    project_text = input(f"{x+1}-project text: ")
    doc.add_paragraph(project_title + " -  "+project_text)
    
# Certificates
doc.add_heading('Certificates', level=2)
num_certs = int(input("How many Certificates you wanna show: "))
for x in range(num_certs):
    certificate = input(f"{x+1}-project title: ")
    doc.add_paragraph(certificate)

# Save the document
doc.save('cv.docx')
